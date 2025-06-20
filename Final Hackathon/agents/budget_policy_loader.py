import json
import pandas as pd
import os
from typing import Dict, Any
from langchain.tools import Tool
from langchain.prompts import PromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.agents import initialize_agent, AgentType
from base_agent import BaseAgent
import config
import PyPDF2
import docx
import re

class BudgetPolicyLoaderAgent(BaseAgent):
    def __init__(self):
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=config.GEMINI_API_KEY
        )
        self.vector_store = None
        self.budget_data = {}
        self.extracted_content = ""
        super().__init__()
    
    def create_tools(self):
        return [
            Tool(
                name="extract_budget_from_file",
                description="Extract budget data from uploaded Excel, PDF, or DOCX file. Input should be a file path.",
                func=self.extract_budget_from_file
            ),
            Tool(
                name="structure_budget_rules",
                description="Structure extracted budget data into JSON format with departments/vendors/categories. Input should be extracted content text.",
                func=self.structure_budget_rules
            ),
            Tool(
                name="create_rag_vectorstore",
                description="Create RAG vector store for budget policy retrieval. Input should be structured content.",
                func=self.create_rag_vectorstore
            )
        ]
    
    def create_prompt(self):
        return PromptTemplate(
            input_variables=["input", "agent_scratchpad", "tools", "tool_names"],
            template="""You are a Budget Policy Loader Agent with RAG capabilities.

Your mission: Extract and structure budget rules from company documents.

Available tools: {tool_names}
{tools}

Use the following format:

Question: the input question you must answer
Thought: you should always think about what to do
Action: the action to take, should be one of [{tool_names}]
Action Input: the input to the action
Observation: the result of the action
... (this Thought/Action/Action Input/Observation can repeat N times)
Thought: I now know the final answer
Final Answer: the final answer to the original input question

Question: {input}
{agent_scratchpad}"""
        )
    
    def extract_budget_from_file(self, file_path: str) -> str:
        """Extract budget data from uploaded file (Excel, PDF, DOCX)"""
        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at {file_path}"
                
            content = ""
            
            if file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
                content = f"Budget Excel Data:\n{df.to_string()}"
                
            elif file_path.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        content += page.extract_text()
                        
            elif file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
            elif file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
                content = f"Budget CSV Data:\n{df.to_string()}"
            else:
                return f"Error: Unsupported file type for {file_path}"
            
            # Store extracted content for processing
            self.extracted_content = content
            return f"Successfully extracted budget content from {file_path}. Content length: {len(content)} characters.\n\nExtracted content preview:\n{content[:500]}..."
            
        except Exception as e:
            return f"Error extracting budget from file: {str(e)}"
    
    def structure_budget_rules(self, extracted_content: str = None) -> str:
        """Structure budget data into JSON format with departments/vendors/categories"""
        try:
            # Use provided content or stored content
            content_to_process = extracted_content or self.extracted_content
            
            if not content_to_process:
                return "Error: No content available to structure. Please extract content from a file first."
            
            # Enhanced parsing logic using regex
            budget_data = {
                "departments": {},
                "vendors": {},
                "categories": {},
                "general_rules": []
            }
            
            lines = content_to_process.split('\n')
            current_section = None
            current_type = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for department sections
                dept_match = re.search(r'(.*?)\s*department.*budget.*[:\$]?\s*[\$]?(\d+(?:,\d+)*)', line, re.IGNORECASE)
                if dept_match:
                    dept_name = dept_match.group(1).strip()
                    budget_amount = dept_match.group(2).replace(',', '')
                    budget_data["departments"][dept_name] = {
                        "total_budget": int(budget_amount),
                        "categories": {},
                        "constraints": []
                    }
                    current_section = dept_name
                    current_type = "departments"
                    continue
                
                # Check for vendor sections
                vendor_match = re.search(r'vendor\s+(\w+).*budget.*[:\$]?\s*[\$]?(\d+(?:,\d+)*)', line, re.IGNORECASE)
                if vendor_match:
                    vendor_name = vendor_match.group(1).strip()
                    budget_amount = vendor_match.group(2).replace(',', '')
                    budget_data["vendors"][vendor_name] = {
                        "total_budget": int(budget_amount),
                        "constraints": []
                    }
                    current_section = vendor_name
                    current_type = "vendors"
                    continue
                
                # Parse line items with amounts
                amount_match = re.search(r'(.*?)[\:\-]\s*[\$]?(\d+(?:,\d+)*)\s*(.*)', line)
                if amount_match and current_section and current_type:
                    item_name = amount_match.group(1).strip().lstrip('-').strip()
                    amount = int(amount_match.group(2).replace(',', ''))
                    constraints_text = amount_match.group(3).strip()
                    
                    constraints = []
                    if constraints_text:
                        # Extract constraints from parentheses or other patterns
                        constraint_matches = re.findall(r'\((.*?)\)|constraint[:\s]+(.*?)(?:\.|$)', constraints_text, re.IGNORECASE)
                        for match in constraint_matches:
                            constraint = match[0] or match[1]
                            if constraint.strip():
                                constraints.append(constraint.strip())
                    
                    if current_type == "departments":
                        budget_data["departments"][current_section]["categories"][item_name] = {
                            "budget": amount,
                            "constraints": constraints
                        }
                    elif current_type == "vendors":
                        budget_data["vendors"][current_section]["categories"] = budget_data["vendors"][current_section].get("categories", {})
                        budget_data["vendors"][current_section]["categories"][item_name] = {
                            "budget": amount,
                            "constraints": constraints
                        }
            
            # Store structured data
            self.budget_data = budget_data
            
            # Save to file
            os.makedirs(config.UPLOAD_PATH, exist_ok=True)
            with open(f"{config.UPLOAD_PATH}/structured_budget.json", 'w') as f:
                json.dump(budget_data, f, indent=2)
            
            return f"Budget rules structured successfully. Found {len(budget_data['departments'])} departments, {len(budget_data['vendors'])} vendors.\n\nStructured data:\n{json.dumps(budget_data, indent=2)}"
                
        except Exception as e:
            return f"Error structuring budget rules: {str(e)}"
    
    def create_rag_vectorstore(self, content: str = None) -> str:
        """Create FAISS vector store for RAG retrieval of budget policies"""
        try:
            # Use provided content, structured data, or extracted content
            if content:
                document_content = content
            elif self.budget_data:
                document_content = json.dumps(self.budget_data, indent=2)
            elif self.extracted_content:
                document_content = self.extracted_content
            else:
                return "Error: No content available to create vector store."
            
            # Split content into chunks for better retrieval
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", " ", ""]
            )
            
            documents = [Document(
                page_content=document_content,
                metadata={"source": "budget_policy", "type": "company_budget"}
            )]
            
            chunks = text_splitter.split_documents(documents)
            
            # Create FAISS vector store
            self.vector_store = FAISS.from_documents(chunks, self.embeddings)
            
            # Save vector store for RAG retrieval
            os.makedirs(config.VECTOR_STORE_PATH, exist_ok=True)
            self.vector_store.save_local(config.VECTOR_STORE_PATH)
            
            return f"RAG vector store created successfully with {len(chunks)} chunks. Vector store saved to {config.VECTOR_STORE_PATH}."
            
        except Exception as e:
            return f"Error creating RAG vector store: {str(e)}"
    
    def get_budget_data(self):
        """Return structured budget data for other agents"""
        return self.budget_data
    
    def process_budget_file(self, file_path: str) -> str:
        """Main method to process a budget file through the complete pipeline"""
        try:
            # Initialize the agent with tools
            tools = self.create_tools()
            agent = initialize_agent(
                tools=tools,
                llm=self.llm,
                agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
                verbose=True,
                handle_parsing_errors=True
            )
            
            # Process the file
            query = f"Extract budget data from file {file_path}, then structure it into JSON format, and finally create a RAG vector store."
            
            result = agent.run(query)
            return result
            
        except Exception as e:
            # Fallback to direct method calls if agent fails
            print(f"Agent execution failed: {e}. Falling back to direct method calls.")
            
            # Direct pipeline execution
            extract_result = self.extract_budget_from_file(file_path)
            if "Error:" in extract_result:
                return extract_result
            
            structure_result = self.structure_budget_rules()
            if "Error:" in structure_result:
                return structure_result
            
            vector_result = self.create_rag_vectorstore()
            
            return f"Budget file processed successfully:\n\n1. {extract_result}\n\n2. {structure_result}\n\n3. {vector_result}"